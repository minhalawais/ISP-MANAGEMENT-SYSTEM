from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SmartOLT_API_Integration_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
WHITE = "FFFFFF"
GREEN = "276749"
AMBER = "7A5A00"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C7D1DC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element

    def make_numbering(abstract_id, num_id, fmt, text, left=540, hanging=270, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        first_num = numbering.find(qn("w:num"))
        if first_num is None:
            numbering.append(abstract)
        else:
            numbering.insert(numbering.index(first_num), abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    make_numbering(900, 900, "bullet", "\u2022", font="Calibri")
    make_numbering(901, 901, "decimal", "%1.")
    return 900, 901


def new_num_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_ids, default=901) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.insert(0, num_pr)
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_label):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_body(doc, text, bold_lead=None, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return p


def add_callout(doc, label, text, kind="info"):
    color = {"info": BLUE, "success": GREEN, "warning": AMBER, "risk": RED}.get(kind, BLUE)
    fill = {"info": LIGHTER_BLUE, "success": "EEF7F0", "warning": "FFF8E6", "risk": "FDECEC"}.get(kind, LIGHTER_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=color)
    body = p.add_run(text)
    set_run_font(body, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_table(doc, headers, rows, widths_dxa, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, heading in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(heading)
        set_run_font(run, size=9.4, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F9FAFB")
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK, bold=(idx == 0))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    doc.settings.odd_and_even_pages_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    add_page_field(fp)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    r = kicker.add_run("TECHNICAL INTEGRATION GUIDE")
    set_run_font(r, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("SmartOLT API Integration")
    set_run_font(r, size=30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("A complete, practical guide for the Management, Employee and Customer Portals")
    set_run_font(r, size=15, color=DARK_BLUE)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(68)
    p_pr = line._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)

    meta = [
        ("Prepared for", "ISP Management App"),
        ("Document purpose", "SmartOLT API capability, architecture and implementation plan"),
        ("Version", "1.0"),
        ("Prepared", "9 August 2026"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    for row, (label, value) in zip(table.rows, meta):
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run_font(p1.add_run(label), size=10, bold=True, color=MUTED)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(value), size=10, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2160, 7200])
    set_table_borders(table, color="E4E7EC", size="4")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run("Designed for business owners, operations teams, support staff and software engineers")
    set_run_font(r, size=9.5, italic=True, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)
    bullet_id, number_abstract_id = create_numbering(doc)
    add_cover(doc)
    page_break(doc)

    add_heading(doc, "How to Use This Guide", 1)
    add_body(doc, "This document explains what SmartOLT provides, how those capabilities should appear in each portal, what information should be stored in the application database, and how the integration should be implemented safely in production.")
    add_callout(doc, "Core recommendation", "Use a hybrid design. Read normal portal data from the local database, synchronize SmartOLT data in bulk in the background, and call SmartOLT live only for network commands and deliberate diagnostics.", "success")

    add_heading(doc, "Guide at a Glance", 2)
    contents = [
        "Executive Summary",
        "SmartOLT in Plain Language",
        "Integration Principles and System Ownership",
        "Management Portal Module",
        "Employee Portal Features",
        "Customer Portal Features",
        "Data Storage and Synchronization",
        "Operational Workflows",
        "Analytics and Insights Dashboard",
        "Security, Permissions and Audit",
        "Reliability, Rate Limits and Failure Handling",
        "Recommended Data Model and Backend Design",
        "Delivery Roadmap and Acceptance Criteria",
    ]
    contents_num_id = new_num_instance(doc, number_abstract_id)
    for item in contents:
        add_list_item(doc, item, contents_num_id)
    add_list_item(doc, "Appendix A. Complete SmartOLT API Capability Catalogue", bullet_id)

    add_callout(doc, "Scope note", "SmartOLT capabilities can vary by OLT vendor, ONU model, configured features, subscription status and API-key restrictions. The application must use capability checks and clear error messages instead of assuming every command works on every device.", "warning")

    page_break(doc)
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "SmartOLT is a cloud platform for managing fiber OLTs and subscriber ONUs/ONTs. Its API can discover equipment, provision new connections, monitor network health, diagnose faults, change subscriber services and perform lifecycle operations such as rebooting, disabling or replacing an ONU.")
    add_body(doc, "The ISP Management App should present SmartOLT as a dedicated Fiber Network module. The management portal receives the full operational and configuration surface; the employee portal receives role-based installation, support and field-service tools; and the customer portal receives a carefully limited self-service view.")
    add_body(doc, "The local PostgreSQL database should hold synchronized network inventory, current states, historical events, customer mappings, analytics data, command jobs and audit records. SmartOLT remains the source of truth for physical network state and applied OLT/ONU configuration.")

    add_heading(doc, "Key Decisions", 2)
    decisions = [
        "Create a dedicated SmartOLT integration rather than configuring the existing generic per-customer monitoring loop.",
        "Never expose the SmartOLT API token to React, the employee portal or the customer portal.",
        "Use bulk endpoints and cached database reads for normal monitoring screens.",
        "Use a durable command queue for provisioning and configuration changes.",
        "Store desired and applied service settings separately to detect configuration drift.",
        "Strengthen customer authentication before exposing network or connected-device information.",
        "Build a dedicated Fiber Analytics dashboard from locally stored time-series snapshots and events.",
    ]
    for text in decisions:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "2. SmartOLT in Plain Language", 1)
    add_body(doc, "An OLT is the provider-side device that serves many fiber customers. An ONU or ONT is the device installed at the customer's premises. SmartOLT communicates with supported OLTs and provides one management interface for subscriber provisioning, monitoring and configuration.")
    terms = [
        ("OLT", "The central fiber access device operated by the ISP."),
        ("ONU/ONT", "The customer-side optical device connected to the fiber network."),
        ("PON port", "An OLT port that serves multiple customers through splitters."),
        ("Zone", "A SmartOLT service region or operational grouping."),
        ("ODB", "An optical distribution box or splitter location."),
        ("VLAN", "A logical network used to separate internet, IPTV, voice or management traffic."),
        ("Speed profile", "The upload or download policy applied to an ONU."),
        ("TR-069", "A remote-management protocol used to inspect and manage compatible customer equipment."),
        ("External ID", "The stable identifier used to connect a SmartOLT ONU with a record in another system."),
    ]
    add_table(doc, ["Term", "Meaning in this integration"], terms, [1800, 7560], 9.7)

    add_heading(doc, "3. Integration Principles and System Ownership", 1)
    add_callout(doc, "Architecture rule", "Portal pages read the local application API. Background workers synchronize SmartOLT. Live SmartOLT calls are reserved for explicit actions and troubleshooting.", "info")
    ownership = [
        ("Customer identity, invoices and packages", "ISP Management App"),
        ("Customer-to-ONU association", "ISP Management App, validated against SmartOLT"),
        ("Physical ONU/OLT state", "SmartOLT"),
        ("Applied network configuration", "SmartOLT"),
        ("Desired network configuration", "ISP Management App"),
        ("Areas and sub-zones", "ISP Management App"),
        ("SmartOLT zones, ODBs, VLANs and profiles", "SmartOLT"),
        ("Mapping between business and network records", "ISP Management App"),
        ("Command, approval and audit history", "ISP Management App"),
        ("Historical analytics", "ISP Management App database"),
    ]
    add_table(doc, ["Information", "Authoritative owner"], ownership, [4680, 4680], 9.5)

    add_heading(doc, "Customer and Service Mapping", 2)
    mappings = [
        ("Company", "SmartOLT account, subdomain and encrypted API token"),
        ("Customer", "Linked SmartOLT ONU/ONT"),
        ("Internet ID", "Business-facing identifier; not automatically the SmartOLT external ID"),
        ("Primary customer package", "Upload and download speed profiles"),
        ("Package add-ons", "IPTV, CATV, VoIP or a configuration preset"),
        ("Area", "SmartOLT zone through a configurable mapping"),
        ("Sub-zone", "SmartOLT ODB/splitter through a configurable mapping"),
        ("Inventory item", "ONU serial/MAC and ONU type"),
        ("Complaint", "Network diagnostic snapshot and incident context"),
        ("Technician task", "Provisioning, relocation, repair or replacement workflow"),
    ]
    add_table(doc, ["Application record", "SmartOLT relationship"], mappings, [2880, 6480], 9.4)
    add_callout(doc, "Important identifier rule", "Create a separate smartolt_customer_link with a dedicated SmartOLT external ID. SmartOLT documents that external IDs are alphanumeric, so an existing internet ID containing punctuation may not be accepted.", "warning")

    add_heading(doc, "4. Management Portal Module", 1)
    add_body(doc, "Add a top-level Fiber Network module. It should feel like an operational NOC workspace, with compact tables, filters, status indicators, guided actions and explicit permissions.")
    management_modules = [
        ("Network Overview", "OLT health, ONU status, outages, signal warnings, provisioning queue and command failures."),
        ("OLT Infrastructure", "OLTs, cards, PON ports, uplinks, temperature, uptime and saved configurations."),
        ("Provisioning", "Unconfigured ONUs, authorization presets, manual authorization, replacement and relocation."),
        ("ONU Management", "Searchable customer-linked ONU inventory, current state, signal, service and topology."),
        ("Live Diagnostics", "On-demand status, signal, full status, running configuration, hosts, reboot and resync."),
        ("Network Catalogue", "ONU types, zones, ODBs, VLANs, speed profiles and configuration presets."),
        ("Subscriber Configuration", "WAN, TR-069, VoIP, VLAN, speed, management IP and service-port settings."),
        ("Ports and Services", "Ethernet/Wi-Fi modes, IPTV, CATV, DHCP controls and security settings."),
        ("Incidents and Alerts", "Outage incidents, signal degradation, stale data and integration failures."),
        ("Fiber Analytics", "Operational, service-quality and business-impact analysis."),
        ("Jobs and Audit", "Command queue, approvals, request/results, actor history and verification state."),
        ("Integration Settings", "Connection, mappings, sync intervals, rate budget and health checks."),
    ]
    add_table(doc, ["Submodule", "Purpose"], management_modules, [2520, 6840], 9.3)

    add_heading(doc, "Management Actions by Risk", 2)
    risks = [
        ("Read-only", "View cached status, signals, graphs, inventory, topology and outages.", "Support, technician, NOC, manager, auditor"),
        ("Low operational", "On-demand refresh, reboot, resync and update location.", "Technician, NOC, manager"),
        ("Configuration", "Speed, VLAN, WAN, TR-069, service ports and presets.", "NOC, manager"),
        ("High impact", "Bulk enable/disable, PON moves and service-wide changes.", "Manager approval"),
        ("Destructive", "Factory reset and deletion.", "Elevated approval, reason and confirmation"),
    ]
    add_table(doc, ["Risk level", "Examples", "Recommended access"], risks, [1700, 4420, 3240], 8.9)

    add_heading(doc, "5. Employee Portal Features", 1)
    add_body(doc, "Add a Fiber Operations section to the employee portal. What an employee can see or execute must be decided on the backend from their role and assigned scope, not only by hiding buttons in React.")
    employee_roles = [
        ("Support", "Cached status, outage, signal health, graph and applied-package verification."),
        ("Technician", "Live diagnostics, provisioning, replacement, relocation, reboot, resync and location updates."),
        ("NOC", "Advanced configuration, profiles, VLANs, WAN, TR-069, ports and subscriber services."),
        ("Recovery agent", "Request suspension/reactivation linked to invoice recovery; approval recommended."),
        ("Manager", "Bulk commands, approvals, network overview and integration health."),
        ("Auditor", "Read-only history, command outcomes, approvals and configuration changes."),
    ]
    add_table(doc, ["Employee role", "SmartOLT features"], employee_roles, [2160, 7200], 9.4)
    add_heading(doc, "Complaint Assistance", 2)
    add_body(doc, "When an employee opens a complaint, the app should attach the latest cached ONU status, signal, OLT/PON, outage state, applied speed profile and last-seen time. This helps distinguish common causes:")
    for text in [
        "Power failure: likely customer power or equipment issue.",
        "LOS: likely fiber loss, cut, connector or splitter issue.",
        "PON outage: shared infrastructure problem affecting multiple subscribers.",
        "Online with poor signal: degradation requiring optical inspection.",
        "Online with correct signal: investigate WAN, Wi-Fi, customer equipment or service configuration.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "6. Customer Portal Features", 1)
    add_body(doc, "Add a My Connection tab that explains network information in customer-friendly language. Normal page loading must use cached local data and show when it was last updated.")
    customer_features = [
        ("Connection status", "Online, offline, fiber signal lost or power issue, with last update time."),
        ("Current package", "Subscribed speed and the speed profile currently applied to the ONU."),
        ("Signal health", "Simple Good, Needs Attention or Critical label; raw optical values may be optional."),
        ("Known outage", "Area/PON incident notice when the customer is part of an active outage."),
        ("Graphs", "Signal and traffic graph images proxied and cached by the backend."),
        ("Connection check", "Rate-limited live status check initiated by the customer."),
        ("Restart device", "Confirmed, cooldown-controlled reboot when supported."),
        ("Connected devices", "Optional TR-069 host list after strong authentication and explicit request."),
        ("Report a problem", "Creates a complaint prefilled with the latest diagnostic context."),
    ]
    add_table(doc, ["Customer feature", "What it does"], customer_features, [2600, 6760], 9.3)
    add_callout(doc, "Security prerequisite", "The current CNIC-only lookup is not strong enough for connected-device or detailed network information. Add OTP/password authentication, a customer session, lookup rate limiting and event logging before enabling SmartOLT customer features.", "risk")
    add_body(doc, "Do not expose VLANs, WAN or PPPoE credentials, running configuration, factory reset, ONU deletion, administrative disable/enable, port modes or SmartOLT credentials to customers.")

    add_heading(doc, "7. Data Storage and Synchronization", 1)
    add_callout(doc, "Decisive rule", "Read from the local database by default; synchronize in bulk in the background; call SmartOLT live only for commands and deliberate diagnostics.", "success")
    strategy = [
        ("OLTs, cards, PONs and uplinks", "Yes", "Scheduled sync"),
        ("ONU inventory and configuration metadata", "Yes", "Initial import plus incremental sync"),
        ("Zones, ODBs, VLANs, types and profiles", "Yes", "Infrequent or manual sync"),
        ("Customer-to-ONU mapping", "Yes", "Updated during provisioning/linking"),
        ("Current ONU status", "Yes", "Bulk poll every 5-7 minutes"),
        ("Optical signals", "Current and history", "Bulk poll every 15-30 minutes"),
        ("PON outages", "Current and history", "Poll every 3-5 minutes"),
        ("OLT uptime and temperature", "Current and history", "Poll every 5-10 minutes"),
        ("GPS coordinates", "Yes", "Daily or after location changes"),
        ("Subscription details", "Yes", "Daily"),
        ("Signal/traffic graph images", "Short cache", "Fetch when requested"),
        ("Full status and running config", "Short diagnostic retention", "On demand only"),
        ("Connected router hosts", "No history or very short cache", "On demand only"),
        ("Commands, approvals and results", "Always", "Recorded for every operation"),
    ]
    add_table(doc, ["Data", "Store locally", "How SmartOLT is called"], strategy, [3600, 2160, 3600], 8.8)

    add_heading(doc, "Current State and History", 2)
    add_body(doc, "Maintain one current-state row per ONU for fast screens, then write event or time-series records only when useful. For example, status history should record state changes instead of repeating the same Online value every few minutes.")
    retention = [
        ("Raw signal samples", "30-90 days"),
        ("Hourly signal aggregates", "12-24 months"),
        ("Daily aggregates", "Long term"),
        ("Status and outage events", "Long term"),
        ("Command and audit history", "According to compliance policy"),
        ("Connected-device data", "Request-scoped or 1-5 minute cache"),
        ("Running configuration", "Encrypted, short diagnostic retention"),
    ]
    add_table(doc, ["Record type", "Recommended retention"], retention, [4680, 4680], 9.4)

    add_heading(doc, "Desired versus Applied Configuration", 2)
    add_body(doc, "Store both what the business system wants and what SmartOLT reports as applied. A package change is complete only after SmartOLT succeeds and verification confirms the expected profiles.")
    add_callout(doc, "Example", "Desired download profile: 100M; applied download profile: 50M; configuration state: Out of sync. This should create a reconciliation alert instead of silently showing the package as successfully applied.", "info")

    add_heading(doc, "8. Operational Workflows", 1)
    workflows = [
        ("New installation", [
            "Create or approve the customer and installation task.",
            "Discover the unconfigured ONU by serial/MAC and OLT/PON.",
            "Match the device with assigned inventory and customer.",
            "Find the best applicable authorization preset.",
            "Authorize and configure the ONU through the command queue.",
            "Link customer, ONU, package, zone and ODB records.",
            "Verify online state, signal and applied speed profiles.",
            "Complete the task and retain the full audit trail.",
        ]),
        ("Package change", [
            "Save the requested package as desired state.",
            "Resolve upload/download profiles and service add-ons.",
            "Queue the SmartOLT changes and prevent conflicting commands.",
            "Execute, verify and update the applied state.",
            "Notify staff if the desired and applied settings do not match.",
        ]),
        ("Complaint diagnosis", [
            "Load cached state and active outage context.",
            "Run live diagnostics only when an employee explicitly requests them.",
            "Attach findings to the complaint and technician task.",
            "Execute approved reboot/resync/configuration actions.",
            "Verify recovery and record resolution evidence.",
        ]),
        ("Suspension and reactivation", [
            "Create an approval-backed command from billing/recovery status.",
            "Disable or enable the ONU using a bulk call where appropriate.",
            "Verify administrative status and record the billing trigger.",
            "Do not delete an ONU as part of routine suspension.",
        ]),
        ("Equipment replacement", [
            "Validate replacement inventory and technician assignment.",
            "Update the serial/MAC or authorize the replacement device.",
            "Reapply the correct preset, profiles and services.",
            "Verify status and return the old device through inventory workflow.",
        ]),
    ]
    for title, steps in workflows:
        add_heading(doc, title, 2)
        number_id = new_num_instance(doc, number_abstract_id)
        for step in steps:
            add_list_item(doc, step, number_id)

    add_heading(doc, "Command Job Lifecycle", 2)
    states = [
        ("Pending", "Accepted and waiting for validation/worker execution."),
        ("Running", "Sent to SmartOLT or actively processing."),
        ("Succeeded", "SmartOLT accepted the command and verification passed."),
        ("Failed", "SmartOLT rejected the operation or verification proved it did not apply."),
        ("Verification required", "The command outcome is uncertain, for example after a timeout."),
        ("Cancelled", "Stopped before execution by an authorized user."),
    ]
    add_table(doc, ["State", "Meaning"], states, [2520, 6840], 9.4)

    add_heading(doc, "9. Analytics and Insights Dashboard", 1)
    add_body(doc, "Create a separate Fiber Analytics dashboard using locally stored snapshots and events. The dashboard should answer operational questions, explain customer impact and support proactive maintenance.")
    analytics = [
        ("Network availability", "Availability by OLT, PON, zone, ODB and service plan."),
        ("Outage performance", "Incident count, affected customers, duration, recovery and MTTR."),
        ("Signal quality", "Warning/critical counts, degradation trends and repeated weak-signal customers."),
        ("PON health", "Capacity, online ratio, average signal and concentration of failures."),
        ("OLT reliability", "Uptime, temperature, card and uplink faults."),
        ("Geographic risk", "Map clusters of outages, LOS and weak optical signals."),
        ("Customer impact", "Complaints and revenue exposure correlated with outages and signal problems."),
        ("Configuration quality", "Package/profile mismatches, unlinked ONUs and failed reconciliation."),
        ("Provisioning performance", "Success rate and average time from discovery to verified service."),
        ("Team effectiveness", "Technician work, repeat faults and resolution outcomes."),
    ]
    add_table(doc, ["Analysis area", "Purpose"], analytics, [2880, 6480], 9.2)
    add_callout(doc, "Data limitation", "SmartOLT's signal and traffic graph endpoints return PNG images. Use them for display, not numerical analysis. Build numerical analytics from bulk status, signal, OLT and PON data stored in the local database.", "warning")

    add_heading(doc, "10. Security, Permissions and Audit", 1)
    security_items = [
        "Encrypt each company's SmartOLT API token at rest and mask it in all responses and logs.",
        "Send the X-Token header only from the Flask backend.",
        "Enforce company isolation and role permissions on every SmartOLT route.",
        "Require step-up approval for bulk, high-impact and destructive actions.",
        "Record actor, company, reason, target ONU, desired change, request, response and verification result.",
        "Never log PPPoE passwords, Wi-Fi passwords, web credentials or complete API tokens.",
        "Use short-lived customer sessions, OTP/password authentication and per-customer rate limits.",
        "Protect connected-device lists and running configurations with stronger access and short retention.",
        "Provide read-only audit access without exposing secrets.",
    ]
    for item in security_items:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "11. Reliability, Rate Limits and Failure Handling", 1)
    add_body(doc, "SmartOLT documents a default budget of up to 1,000 calls per hour and requires responsible caching, bulk requests and randomized polling intervals. The integration should track consumption per company and slow nonessential jobs before reaching the limit.")
    safeguards = [
        ("Bulk polling", "Use all-ONU endpoints rather than one request per customer."),
        ("Jitter", "Spread scheduled calls across each interval to avoid synchronized spikes."),
        ("Caching", "Serve normal screens from local current-state records."),
        ("Incremental sync", "Use pagination, selected fields and updated_since after the initial ONU import."),
        ("429 handling", "Honor Retry-After and back off without creating a request storm."),
        ("Retries", "Retry safe reads; do not blindly retry state-changing commands."),
        ("Circuit breaker", "Pause repeated calls when SmartOLT or an OLT is failing."),
        ("Stale data", "Continue serving cached records and label them stale; do not invent offline events."),
        ("Uncertain command", "Mark verification required rather than assuming success or failure."),
        ("Conflict control", "Allow only one incompatible command at a time for the same ONU."),
    ]
    add_table(doc, ["Control", "Required behavior"], safeguards, [2520, 6840], 9.3)
    add_callout(doc, "Existing implementation warning", "The current generic monitoring service loops through active customers and fetches each metric separately. Do not use that strategy for SmartOLT. Implement a dedicated SmartOLT bulk synchronization service.", "risk")

    add_heading(doc, "12. Recommended Data Model and Backend Design", 1)
    tables = [
        ("smartolt_connections", "Company subdomain, encrypted token, restrictions, health and rate state."),
        ("smartolt_customer_links", "Customer-to-ONU relationship, external ID, serial/MAC and link status."),
        ("smartolt_olts", "Synchronized OLT inventory and current health."),
        ("smartolt_olt_cards / pon_ports / uplinks", "Infrastructure topology and current state."),
        ("smartolt_onus", "Normalized ONU inventory, topology, capabilities and applied configuration summary."),
        ("smartolt_onu_current_state", "Fast current status, signal, administrative and CATV state."),
        ("smartolt_status_events", "State transitions and incident durations."),
        ("smartolt_signal_samples", "Optical history for trend and degradation analysis."),
        ("smartolt_outage_incidents", "PON/zone outages, impact and recovery."),
        ("smartolt_catalog_*", "ONU types, zones, ODBs, VLANs, profiles and presets."),
        ("smartolt_mapping_rules", "Area/zone, sub-zone/ODB, package/profile and inventory/type mappings."),
        ("smartolt_command_jobs", "Durable requested operations and lifecycle state."),
        ("smartolt_command_attempts", "Every outbound request, sanitized response and retry/verification record."),
        ("smartolt_sync_runs", "Endpoint, timing, item counts, API usage and errors."),
    ]
    add_table(doc, ["Table/group", "Purpose"], tables, [3420, 5940], 9.0)

    add_heading(doc, "Backend Components", 2)
    components = [
        ("SmartOLTClient", "Typed HTTP client, X-Token authentication, timeout, error normalization and rate handling."),
        ("SmartOLTSyncService", "Bulk inventory/status/signal/topology synchronization and change detection."),
        ("SmartOLTCommandService", "Validation, approvals, command queue, execution and verification."),
        ("SmartOLTMappingService", "Business-to-network mapping and reconciliation."),
        ("SmartOLTAnalyticsService", "Aggregation, incidents, trends and business correlations."),
        ("SmartOLTAccessPolicy", "Server-side role, company, zone and action authorization."),
    ]
    add_table(doc, ["Component", "Responsibility"], components, [2880, 6480], 9.2)

    add_heading(doc, "13. Delivery Roadmap and Acceptance Criteria", 1)
    phases = [
        ("Phase 1 - Foundation", "Connection settings, encrypted credentials, typed client, mappings, bulk inventory/status/signal sync and audit framework."),
        ("Phase 2 - Management visibility", "OLT/ONU screens, outages, topology, customer linking and cached diagnostics."),
        ("Phase 3 - Provisioning", "Unconfigured queue, presets, authorization, installation workflow and verification."),
        ("Phase 4 - Employee operations", "Complaint diagnostics, technician actions, replacement and relocation."),
        ("Phase 5 - Service control", "Speed, enable/disable, VLAN, WAN, TR-069, IPTV/CATV/VoIP and port controls."),
        ("Phase 6 - Customer self-service", "Secure authentication, connection view, check, reboot and complaint creation."),
        ("Phase 7 - Analytics", "History, incidents, trends, alerts and Fiber Analytics dashboard."),
        ("Phase 8 - Hardening", "Load tests, rate-budget simulation, failure recovery, permissions review and production runbooks."),
    ]
    add_table(doc, ["Phase", "Deliverable"], phases, [2520, 6840], 9.2)

    add_heading(doc, "Production Acceptance Checklist", 2)
    acceptance = [
        "No frontend can obtain the SmartOLT API token.",
        "Normal portal loads do not call SmartOLT directly.",
        "Bulk synchronization replaces per-customer monitoring calls.",
        "All records and actions enforce company isolation.",
        "High-impact commands require the intended role and approval.",
        "Every command has a durable job, audit trail and verification result.",
        "Stale cached data is displayed clearly during an upstream outage.",
        "Rate-limit and Retry-After behavior is tested.",
        "Desired/applied mismatch detection works for packages and services.",
        "Customer SmartOLT features require stronger authentication than CNIC-only lookup.",
        "Analytics operate from local numeric data rather than graph-image parsing.",
        "Backups, retention, monitoring and operational runbooks are documented.",
    ]
    for item in acceptance:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Appendix A. Complete SmartOLT API Capability Catalogue", 1)
    add_body(doc, "This catalogue lists the public SmartOLT API functions identified in the official documentation and explains their intended purpose in the ISP Management App. Exact parameters and hardware support must be validated against the live SmartOLT documentation during implementation.")

    endpoint_groups = [
        ("A1. OLT and System Inventory", [
            ("Get OLTs list", "Synchronize OLT names, addresses, management ports and unique IDs."),
            ("Get OLTs uptime and environment temperature", "Monitor availability, restart history and overheating risk."),
            ("Get OLT cards details", "Inspect slots, card types, software versions, roles and health."),
            ("Get OLT PON ports details", "Inspect port state, capacity, ONU counts, average signal and TX power."),
            ("Get OLT outage PONs details", "Retrieve grouped active LOS, power and unknown PON outages for NOC triage."),
            ("Get OLT uplink ports details", "Inspect uplink state, MTU, negotiation, wavelength and tagged VLANs."),
            ("Save OLTs config", "Persist OLT configuration after controlled changes."),
            ("Get billing details", "Monitor SmartOLT OLT subscription status and expiration."),
        ]),
        ("A2. Network Catalogue", [
            ("Get ONU types list", "Synchronize all supported ONU models and capabilities."),
            ("Get ONU types by PON type", "Filter compatible models for GPON or EPON workflows."),
            ("Get ONU type image", "Display the hardware model image through a backend proxy."),
            ("Add ONU type", "Register a model with ports, Wi-Fi, VoIP, CATV and routing capabilities."),
            ("Get zones list / Add zone", "Read or create SmartOLT service regions."),
            ("Get ODBs list / Add ODB", "Read or create splitter/distribution-box topology."),
            ("Get speed profiles list", "Synchronize available upload/download policy profiles."),
            ("Get VLANs list / Add VLAN", "Read or create internet, IPTV and management/VoIP VLANs."),
        ]),
        ("A3. Discovery and Bulk Monitoring", [
            ("Get all unconfigured ONUs", "Discover equipment awaiting authorization across all OLTs."),
            ("Get unconfigured ONUs by OLT", "Discover pending equipment on one OLT."),
            ("Get all ONU statuses", "Bulk cache Online, Power fail, LOS and Offline states."),
            ("Get all ONU administrative statuses", "Bulk cache Enabled/Disabled state."),
            ("Get all ONU CATV statuses", "Bulk cache CATV service state and unsupported devices."),
            ("Get all ONU signals", "Bulk cache optical signal classification and current values."),
            ("Get all ONU details", "Import and incrementally synchronize ONU inventory/configuration metadata."),
            ("Get all ONU GPS coordinates", "Synchronize subscriber locations for topology and map analysis."),
        ]),
        ("A4. ONU Lookup and Live Diagnostics", [
            ("Get ONU status by external ID", "Perform a fresh status check for one ONU during troubleshooting."),
            ("Get ONU signal by external ID", "Perform a fresh optical reading for one ONU."),
            ("Get ONU details by external ID", "Retrieve identity and OLT/board/port placement."),
            ("Get ONUs details by SN", "Find ONU records matching a serial/MAC."),
            ("Get ONU full status info", "Live status history, optics, interfaces, WAN and MAC-table diagnostics."),
            ("Get ONU running config", "Retrieve the active ONU configuration from the OLT."),
            ("Get ONU signal graph", "Return hourly to yearly signal graph PNG images."),
            ("Get ONU traffic graph", "Return hourly to yearly traffic graph PNG images."),
            ("Get ONU speed profiles", "Verify currently assigned upload and download profiles."),
            ("Get ONU router hosts", "Retrieve connected LAN hosts live through TR-069."),
        ]),
        ("A5. Authorization and Provisioning", [
            ("Get authorization presets by OLT", "List reusable authorization rules for an OLT."),
            ("Get applicable authorization presets", "Score presets against an ONU's serial, board and port."),
            ("Authorize ONU using preset", "Provision an ONU through a controlled reusable template."),
            ("Authorize ONU manually", "Provision with explicit OLT, PON, VLAN, type, zone, mode and service settings."),
            ("Move ONU", "Relocate an ONU to another OLT, board or PON port."),
            ("Update ONU PON channel", "Change GPON/XGPON/XGSPON/EPON/10G-EPON channel."),
            ("Update ONU SN/MAC", "Replace or correct the subscriber device serial/MAC."),
            ("Update ONU type", "Change the assigned hardware model."),
            ("Update ONU custom profile", "Apply a vendor/model-specific custom profile."),
        ]),
        ("A6. Configuration Presets", [
            ("Get ONU configuration presets", "List reusable configuration packages visible to the API key."),
            ("Apply configuration preset to one ONU", "Apply and push a reusable configuration, with rollback on failure."),
            ("Remove configuration preset from one ONU", "Remove a previously applied preset."),
            ("Apply configuration preset to multiple ONUs", "Bulk-apply a preset to up to 50 ONU IDs."),
            ("Remove configuration preset from multiple ONUs", "Bulk-remove a preset from up to 50 ONU IDs."),
            ("Check configuration preset task status", "Track the progress and result of a bulk preset task."),
        ]),
        ("A7. Identity, Location and VLAN", [
            ("Update ONU location details", "Update subscriber name/address, zone, ODB and location metadata."),
            ("Update ONU external ID", "Change the CRM/ERP integration identifier."),
            ("Update external ID by board/port/ONU number", "Link a device when topology coordinates are known."),
            ("Update external ID by SN", "Link a device when its serial/MAC is known."),
            ("Update ONU attached VLANs", "Change the VLANs available to the ONU."),
            ("Update ONU main VLAN-ID", "Change the primary subscriber service VLAN."),
            ("Update ONU mode", "Change supported bridging/routing operating mode."),
        ]),
        ("A8. Management IP, TR-069 and Voice", [
            ("Set management IP Inactive", "Disable the ONU management IP service."),
            ("Set management IP Static", "Apply a fixed management IP configuration."),
            ("Set management IP DHCP", "Obtain management addressing dynamically."),
            ("Enable / Disable TR-069", "Turn remote CPE-management integration on or off."),
            ("Set VoIP mode Enabled / Disabled", "Control the ONU voice-service mode."),
            ("Enable / Disable VoIP port", "Control an individual physical voice port."),
        ]),
        ("A9. WAN Configuration", [
            ("Set WAN mode to Setup via ONU webpage", "Leave WAN setup for local/manual device configuration."),
            ("Set WAN mode to DHCP", "Configure dynamic WAN addressing."),
            ("Set WAN mode to Static IP", "Configure fixed WAN addressing and gateway details."),
            ("Set WAN mode to PPPoE", "Configure subscriber PPPoE service credentials."),
            ("Set WAN configuration method", "Select the supported method used to push WAN settings."),
            ("Set WAN IP version", "Select supported IPv4/IPv6 behavior."),
        ]),
        ("A10. Subscriber Security and Access", [
            ("Update maximum MAC learning", "Limit the number of learned downstream MAC addresses."),
            ("Enable / Disable IP DHCP snooping", "Control protection against unauthorized DHCP behavior."),
            ("Enable / Disable DHCP Option 82", "Insert or suppress subscriber circuit-identification information."),
            ("Enable / Disable IP source guard", "Restrict traffic to permitted source addresses."),
            ("Enable / Disable remote WAN access", "Control remote access to the ONU's WAN IP."),
            ("Change ONU web user and password", "Update local ONU administration credentials."),
        ]),
        ("A11. Speed and Service Ports", [
            ("Update ONU speed profiles", "Change one subscriber's upload/download profiles."),
            ("Update multiple ONU speed profiles", "Change profiles for a controlled batch of ONU IDs."),
            ("Update ONU service port", "Modify service-port configuration used by subscriber traffic."),
        ]),
        ("A12. Ethernet and Wi-Fi Ports", [
            ("Ethernet/Wi-Fi port: LAN", "Provide normal subscriber LAN service."),
            ("Ethernet/Wi-Fi port: IPTV", "Assign multicast television service."),
            ("Ethernet/Wi-Fi port: Access", "Use one untagged access VLAN."),
            ("Ethernet/Wi-Fi port: Hybrid", "Combine a primary untagged VLAN with allowed tagged VLANs."),
            ("Ethernet/Wi-Fi port: Trunk", "Carry multiple tagged VLANs."),
            ("Ethernet/Wi-Fi port: Transparent", "Pass traffic transparently where supported."),
            ("Ethernet/Wi-Fi port: Shutdown", "Administratively disable an individual port."),
        ]),
        ("A13. IPTV and CATV", [
            ("Enable / Disable IPTV", "Turn subscriber IPTV service on or off."),
            ("Enable / Disable CATV", "Turn CATV on or off for one ONU."),
            ("Enable / Disable CATV for multiple ONUs", "Control CATV in batches of up to 50 ONU IDs."),
        ]),
        ("A14. ONU Lifecycle", [
            ("Reboot ONU", "Restart the subscriber ONU during controlled troubleshooting."),
            ("Resync ONU config", "Push the stored configuration to the OLT/ONU again."),
            ("Restore ONU factory defaults", "Erase customer configuration and return the ONU to factory state."),
            ("Disable / Enable ONU", "Suspend or restore one subscriber's ONU administratively."),
            ("Disable / Enable multiple ONUs", "Suspend or restore a batch of up to 50 ONU IDs."),
            ("Delete ONU", "Remove the ONU from SmartOLT after approval and retention checks."),
        ]),
    ]

    for heading, rows in endpoint_groups:
        add_heading(doc, heading, 2)
        add_table(doc, ["SmartOLT function", "Purpose in the ISP Management App"], rows, [3600, 5760], 8.8)

    add_heading(doc, "Appendix B. Recommended API Usage Rules", 1)
    rules = [
        ("All ONU statuses", "Every 5-7 minutes", "Bulk call; cache between runs"),
        ("All ONU signals", "Every 15-30 minutes", "Bulk call; retain controlled history"),
        ("PON outages", "Every 3-5 minutes", "Create/update local incidents"),
        ("OLT health", "Every 5-10 minutes", "Store current state and selected history"),
        ("ONU details", "Initial plus incremental", "Paginate; use fields and updated_since"),
        ("GPS coordinates", "Daily or after changes", "SmartOLT documents a strict low-frequency limit"),
        ("Live status/signal/full status", "Explicit troubleshooting only", "Never use in monitoring loops"),
        ("Router hosts", "Explicit troubleshooting only", "One in-flight request; short/no cache"),
        ("Bulk mutations", "When business workflow requires", "Batch up to 50 ONU IDs"),
    ]
    add_table(doc, ["API family", "Starting frequency", "Rule"], rules, [2880, 2520, 3960], 8.9)

    add_heading(doc, "Sources and Implementation References", 1)
    add_body(doc, "The SmartOLT API is a live vendor service and may change. Engineers should confirm endpoint parameters, supported fields and current restrictions immediately before implementation.")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Official SmartOLT API documentation: "), bold=True, color=INK)
    add_hyperlink(p, "Open the official API guide", "https://api.smartolt.com/")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("SmartOLT product information: "), bold=True, color=INK)
    add_hyperlink(p, "Open the SmartOLT website", "https://www.smartolt.com/")
    add_body(doc, "Application references reviewed: src/pages/CustomerPortalPage.tsx; src/pages/EmployeePortal.tsx; src/pages/CustomerDetailPage.tsx; src/components/forms/APIConnectionForm.tsx; api/app/services/monitoring_service.py; api/app/models.py.", color=MUTED)
    doc.core_properties.title = "SmartOLT API Integration Guide"
    doc.core_properties.subject = "Production integration guide for the ISP Management App"
    doc.core_properties.author = "ISP Management App Engineering"
    doc.core_properties.keywords = "SmartOLT, OLT, ONU, ISP, API integration, fiber network"
    doc.core_properties.comments = "Version 1.0"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
